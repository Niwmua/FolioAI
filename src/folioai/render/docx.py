"""DOCX via python-docx (brief §14).

Headings use Word's built-in ``Heading N`` styles rather than bold body text, so Word's own
table-of-contents field finds them. A DOCX whose headings are just large text looks right
and behaves wrong, which is the worst of both.
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from ..errors import RenderError
from ..ir import Block, Document
from ..logging_setup import get_logger
from .base import RenderContext, document_metadata, is_rtl, iter_chapters

log = get_logger(__name__)

_INLINE_RE = re.compile(r"(\*\*.+?\*\*|\*.+?\*|`.+?`)", re.DOTALL)


def _require_docx() -> Any:
    try:
        import docx
    except ImportError as exc:
        raise RenderError(
            "DOCX output needs the 'python-docx' package, which is not installed.",
            remedy="Install the render extra: uv sync --extra render",
            context={"format": "docx"},
        ) from exc
    return docx


def _add_runs(paragraph: Any, text: str) -> None:
    """Add text to a paragraph, converting the IR's inline subset into real runs."""
    for piece in _INLINE_RE.split(text):
        if not piece:
            continue
        if piece.startswith("**") and piece.endswith("**"):
            paragraph.add_run(piece[2:-2]).bold = True
        elif piece.startswith("*") and piece.endswith("*"):
            paragraph.add_run(piece[1:-1]).italic = True
        elif piece.startswith("`") and piece.endswith("`"):
            run = paragraph.add_run(piece[1:-1])
            run.font.name = "Consolas"
        else:
            paragraph.add_run(piece)


def _add_block(doc: Any, block: Block, context: RenderContext) -> None:
    if block.kind == "heading":
        doc.add_heading(block.text, level=min(max(block.level or 1, 1), 4))
        return
    if block.kind == "scene_break":
        paragraph = doc.add_paragraph("* * *")
        paragraph.alignment = 1  # centred
        return
    if block.kind == "page_break":
        doc.add_page_break()
        return

    if context.bilingual:
        source = context.source_text(block.id)
        if source:
            paragraph = doc.add_paragraph(style="Quote")
            run = paragraph.add_run(source)
            run.italic = True
            run.font.size = None

    style = {
        "blockquote": "Quote",
        "list_item": "List Bullet",
        "footnote": "Caption",
        "figure_caption": "Caption",
    }.get(block.kind)
    try:
        paragraph = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    except KeyError:
        # A template without that built-in style: body text is better than a crash.
        log.debug("docx_style_missing", style=style, block=block.id)
        paragraph = doc.add_paragraph()
    _add_runs(paragraph, block.text)


def render_docx(document: Document, path: Path, context: RenderContext | None = None) -> Path:
    """Write a DOCX with styled headings so Word's TOC field works."""
    docx = _require_docx()
    context = context or RenderContext()
    meta = document_metadata(document, context)

    doc = docx.Document()
    doc.core_properties.title = meta["title"]
    doc.core_properties.language = meta["language"]
    if document.author:
        doc.core_properties.author = document.author
    if "translator" in meta:
        doc.core_properties.comments = meta["translator"]

    if is_rtl(meta["language"]):
        # python-docx has no first-class RTL switch; the note travels with the file so a
        # human setting it in Word knows it was expected rather than forgotten.
        doc.core_properties.keywords = "rtl"

    for chapter in iter_chapters(document):
        for block in chapter.blocks:
            _add_block(doc, block, context)

    path.parent.mkdir(parents=True, exist_ok=True)
    try:
        doc.save(str(path))
    except OSError as exc:
        raise RenderError(
            f"Could not write the DOCX to {path}: {exc}",
            remedy="Check the directory is writable and the file is not open in Word.",
            context={"path": str(path)},
        ) from exc
    log.info("docx_written", path=str(path))
    return path
